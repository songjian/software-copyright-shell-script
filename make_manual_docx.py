#!/usr/bin/env python
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Cm
import os
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式

path = sys.argv[1] + sys.argv[2] + '使用说明书.docx'
os.system("touch %s" %path)
document = Document('manual-temp.docx')
document.styles['Normal'].font.name = u'Times New Roman' #设置西文字体
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体使用字体2->宋体
document.sections[0].header.paragraphs[0].text=sys.argv[2] + '-使用说明书'
t1=document.add_heading(sys.argv[2], 1)
t1.alignment=WD_ALIGN_PARAGRAPH.CENTER
t2=document.add_heading('使用说明书', 1)
t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
document.save(path) 
